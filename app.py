import streamlit as st
import pytesseract
import pandas as pd

from PIL import Image
from pdf2image import convert_from_bytes
from deep_translator import GoogleTranslator
from langdetect import detect

from docx import Document
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet

from io import BytesIO
import zipfile

##########################################################
# CONFIG
##########################################################

st.set_page_config(
    page_title="Professional OCR Translator",
    page_icon="📄",
    layout="wide"
)

# WINDOWS USERS
# Uncomment and edit path if needed

# pytesseract.pytesseract.tesseract_cmd = (
#     r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# )

##########################################################
# FUNCTIONS
##########################################################

def ocr_image(image):

    data = pytesseract.image_to_data(
        image,
        output_type=pytesseract.Output.DATAFRAME
    )

    confidence = data["conf"].mean()

    text = pytesseract.image_to_string(image)

    return text, confidence


def ocr_pdf(pdf_bytes):

    pages = convert_from_bytes(pdf_bytes)

    full_text = ""
    scores = []

    for page in pages:

        txt, conf = ocr_image(page)

        full_text += txt + "\n\n"

        scores.append(conf)

    return full_text, sum(scores) / len(scores)


def translate_text(text, target_lang):

    try:

        source_lang = detect(text)

    except:

        source_lang = "auto"

    chunk_size = 4500

    chunks = [
        text[i:i+chunk_size]
        for i in range(
            0,
            len(text),
            chunk_size
        )
    ]

    translated_chunks = []

    for chunk in chunks:

        translated = GoogleTranslator(
            source='auto',
            target=target_lang
        ).translate(chunk)

        translated_chunks.append(translated)

    return source_lang, "\n\n".join(
        translated_chunks
    )


##########################################################
# EXPORT FUNCTIONS
##########################################################

def create_word(original, translated):

    doc = Document()

    doc.add_heading(
        "OCR Translation Report",
        0
    )

    doc.add_heading(
        "Original Text",
        level=1
    )

    doc.add_paragraph(original)

    doc.add_heading(
        "Translated Text",
        level=1
    )

    doc.add_paragraph(translated)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


def create_pdf(original, translated):

    buffer = BytesIO()

    pdf = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    content = []

    content.append(
        Paragraph(
            "OCR Translation Report",
            styles["Title"]
        )
    )

    content.append(
        Paragraph(
            "Original Text",
            styles["Heading1"]
        )
    )

    content.append(
        Paragraph(
            original.replace("\n","<br/>"),
            styles["BodyText"]
        )
    )

    content.append(PageBreak())

    content.append(
        Paragraph(
            "Translated Text",
            styles["Heading1"]
        )
    )

    content.append(
        Paragraph(
            translated.replace("\n","<br/>"),
            styles["BodyText"]
        )
    )

    pdf.build(content)

    buffer.seek(0)

    return buffer


def create_excel(original, translated):

    df = pd.DataFrame(
        {
            "Original":[original],
            "Translated":[translated]
        }
    )

    output = BytesIO()

    with pd.ExcelWriter(
        output,
        engine="openpyxl"
    ) as writer:

        df.to_excel(
            writer,
            index=False
        )

    output.seek(0)

    return output


def create_txt(original, translated):

    buffer = BytesIO()

    content = f"""
ORIGINAL TEXT
======================

{original}

TRANSLATED TEXT
======================

{translated}
"""

    buffer.write(content.encode())

    buffer.seek(0)

    return buffer

##########################################################
# UI
##########################################################

st.title("📄 Professional OCR Translator")

st.markdown(
"""
Upload images or PDFs.
Extract text.
Translate.
Export to multiple formats.
"""
)

target_language = st.selectbox(
    "Translate To",
    [
        "en",
        "fr",
        "de",
        "es",
        "pt",
        "it",
        "nl",
        "pl",
        "ru",
        "zh-CN",
        "ja"
    ]
)

files = st.file_uploader(
    "Upload Files",
    type=[
        "jpg",
        "jpeg",
        "png",
        "tif",
        "tiff",
        "pdf"
    ],
    accept_multiple_files=True
)

if files:

    zip_buffer = BytesIO()

    with zipfile.ZipFile(
        zip_buffer,
        "w"
    ) as zip_file:

        for file in files:

            st.divider()

            st.subheader(file.name)

            if file.type == "application/pdf":

                text, confidence = ocr_pdf(
                    file.read()
                )

            else:

                image = Image.open(file)

                st.image(
                    image,
                    width=500
                )

                text, confidence = ocr_image(
                    image
                )

            source_lang, translated = (
                translate_text(
                    text,
                    target_language
                )
            )

            st.metric(
                "OCR Confidence",
                f"{confidence:.1f}%"
            )

            st.metric(
                "Detected Language",
                source_lang
            )

            col1, col2 = st.columns(2)

            with col1:

                st.subheader("OCR")

                st.text_area(
                    "",
                    text,
                    height=400,
                    key=f"ocr_{file.name}"
                )

            with col2:

                st.subheader("Translation")

                st.text_area(
                    "",
                    translated,
                    height=400,
                    key=f"trans_{file.name}"
                )

            word_file = create_word(
                text,
                translated
            )

            pdf_file = create_pdf(
                text,
                translated
            )

            excel_file = create_excel(
                text,
                translated
            )

            txt_file = create_txt(
                text,
                translated
            )

            base = file.name.split(".")[0]

            zip_file.writestr(
                f"{base}.docx",
                word_file.read()
            )

            zip_file.writestr(
                f"{base}.pdf",
                pdf_file.read()
            )

            zip_file.writestr(
                f"{base}.xlsx",
                excel_file.read()
            )

            zip_file.writestr(
                f"{base}.txt",
                txt_file.read()
            )

    zip_buffer.seek(0)

    st.download_button(
        "📦 Download All Results",
        zip_buffer,
        file_name="OCR_Translations.zip",
        mime="application/zip"
    )

st.success(
    "Supports historical documents, scanned books, municipal records, archives, and handwritten-enhanced OCR workflows."
)